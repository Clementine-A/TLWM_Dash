"""
Extraction du contenu brut de tous les rapports Word DNMF TOGO
Affiche le texte de chaque fichier pour identifier la structure réelle des données
"""
import os
import docx

BASE = r"c:\Users\ATD\Desktop\cle\TLWM\dash"

def extract_file(path, label):
    print(f"\n{'='*80}")
    print(f"FICHIER: {label}")
    print(f"{'='*80}")
    try:
        doc = docx.Document(path)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[P{i:03d}] {para.text.strip()}")
        # Tables
        for t_idx, table in enumerate(doc.tables):
            print(f"\n--- TABLE {t_idx} ---")
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    print(" | ".join(cells))
    except Exception as e:
        print(f"ERREUR: {e}")

# Lire quelques fichiers représentatifs
files_to_check = [
    (os.path.join(BASE, "2024", "Rapport Mensuel DNMF TOGO Janvier.docx"), "JANVIER 2024"),
    (os.path.join(BASE, "2024", "Rapport Mensuel DNMF TOGO JUIN 2024.docx"), "JUIN 2024"),
    (os.path.join(BASE, "2024", "Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx"), "DECEMBRE 2024"),
    (os.path.join(BASE, "2025", "DAMF RAPPORT NATIONAL TOGO Jan 2025.docx"), "JANVIER 2025"),
    (os.path.join(BASE, "2025", "DNMF TOGO- RAPPORT MENSUEL MOIS DE JUIN 2025.docx"), "JUIN 2025"),
]

for path, label in files_to_check:
    extract_file(path, label)
