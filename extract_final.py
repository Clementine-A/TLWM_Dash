"""
Extracteur FINAL - DNMF TOGO 2024 & 2025
Structure 2024: Tables cumulatives (rows=mois, col=valeurs)
Structure 2025 Jan/Fev: grand tableau multi-colonnes (idem 2025 mars+)
Structure 2025 Mars+: tableau annuel (rows=indicateurs, cols=mois)
"""
import os, re, json
import unicodedata
import docx

BASE = r"c:\Users\ATD\Desktop\cle\TLWM\dash"

MONTH_NAMES_DISPLAY = [
    'Janvier','F\u00e9vrier','Mars','Avril','Mai','Juin',
    'Juillet','Ao\u00fbt','Septembre','Octobre','Novembre','D\u00e9cembre'
]

def norm(s):
    """Normalize text: lowercase, remove accents"""
    s = unicodedata.normalize('NFD', str(s).lower())
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn').strip()

MONTH_ID = {
    'janvier': 0, 'janv': 0, 'jan': 0,
    'fevrier': 1, 'fev': 1, 'feb': 1,
    'mars': 2,
    'avril': 3, 'avr': 3,
    'mai': 4,
    'juin': 5,
    'juillet': 6, 'juil': 6, 'juill': 6,
    'aout': 7, 'aou': 7,
    'septembre': 8, 'sept': 8, 'sep': 8,
    'octobre': 9, 'oct': 9,
    'novembre': 10, 'nov': 10,
    'decembre': 11, 'dec': 11,
}

def safe_int(s):
    """Convert cell text to int, handling embedded text like '7445 (6273...)'"""
    s = str(s).strip()
    # Take only the first number found
    m = re.search(r'\d[\d\s]*', s)
    if m:
        try:
            return int(m.group(0).replace(' ', ''))
        except:
            pass
    return 0

# ─── 2024: cumulative tables (rows=months, cols=values) ──────
def extract_2024(doc, target_month_idx):
    data = {}

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # Row 0 = section header, Row 1 = column names, Row 2+ = month rows
        r0 = norm(table.rows[0].cells[0].text)
        r1 = norm(table.rows[1].cells[0].text) if len(table.rows) > 1 else ''

        # === Table 0: Ressources humaines (1 data row) ===
        if t_idx == 0 and ('predicateur' in r0 or 'nombre' in r0):
            if len(table.rows) >= 2:
                row = table.rows[1]
                data['predicateurs'] = safe_int(row.cells[0].text) if len(row.cells) > 0 else 0
                data['eleves_inscrits'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                data['miss_nationaux'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                data['miss_internationaux'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0
                data['predicateurs_utilises'] = safe_int(row.cells[4].text) if len(row.cells) > 4 else 0
            continue

        # === Tables with rows=months: need row1 == 'mois' ===
        if r1 != 'mois':
            # Check assemblées count
            if 'assemblee' in r0 and 'nombre' in r0:
                for row in table.rows[1:]:
                    val = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                    if val > 0:
                        data['assemblees_count'] = val
                        break
            continue

        # Header identifies section
        section = r0

        # Parse month rows
        for row in table.rows[2:]:
            if not row.cells:
                continue
            month_label = norm(row.cells[0].text)
            # Try exact match then prefix
            mi = MONTH_ID.get(month_label)
            if mi is None:
                for k, v in MONTH_ID.items():
                    if month_label.startswith(k):
                        mi = v
                        break
            if mi != target_month_idx:
                continue

            # Séminaires / Urgences
            if 'seminaire' in section or 'urgence' in section:
                data['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                data['assistance'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                data['sauves'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

            # Suivi des sauvés
            elif 'suivi' in section and 'sauve' in section:
                if not data.get('sauves'):
                    data['sauves'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                data['ajoutes'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0

            # Répartition
            elif 'repartition' in section:
                data['sem_assemblees'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                data['sem_hors'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0
                if not data.get('sem_total'):
                    data['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
            break

    return data

# ─── 2025 Mars+: rows=indicateurs, cols=months ───────────────
def find_month_col(header_row, month_idx):
    for i, cell in enumerate(header_row.cells):
        tok = norm(cell.text).split()[0] if norm(cell.text).split() else ''
        if MONTH_ID.get(tok) == month_idx:
            return i
    return -1

def row_val(table, kw, col):
    for row in table.rows:
        if kw in norm(row.cells[0].text):
            return safe_int(row.cells[col].text) if col < len(row.cells) else 0
    return 0

def extract_2025_multi(doc, month_idx):
    data = {}
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h0 = norm(table.rows[0].cells[0].text)
        hdr = table.rows[1]
        col = find_month_col(hdr, month_idx)
        if col < 0:
            continue

        if 'sce' in h0 or 'etat' in h0:
            data['assemblees_count'] = row_val(table, 'assembl', col)
            data['membres'] = row_val(table, 'membres', col)
            data['districts'] = row_val(table, 'district', col)
        elif 'semin' in h0 or ('sem' in h0 and h0.startswith('sem')):
            data['sem_assemblees'] = row_val(table, 'assembl', col)
            data['sem_hors'] = row_val(table, 'hors', col)
            data['sem_total'] = row_val(table, 'total', col)
            data['assistance'] = row_val(table, 'assistance', col)
            data['sauves'] = row_val(table, 'sauv', col)
            data['ajoutes'] = row_val(table, 'ajout', col)
            data['invites'] = row_val(table, 'invit', col)
            data['temoignages'] = row_val(table, 'temoignage', col)
            data['predicateurs_utilises'] = row_val(table, 'predicateurs utilis', col)
        elif 'spgfm' in h0:
            data['predicateurs'] = row_val(table, 'predicateurs', col)
            data['pasteurs'] = row_val(table, 'pasteurs', col)
            data['miss_nationaux'] = row_val(table, 'national', col)
            data['miss_internationaux'] = row_val(table, 'international', col)
        elif 'sfa' in h0:
            data['eleves_inscrits'] = row_val(table, 'inscrits', col)
            data['eleves_actuels'] = row_val(table, 'actuels', col)
    return data

def get_diffs_persps(doc):
    paras = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]
    diffs, persps = [], []
    mode = None
    for p in paras:
        pn = norm(p)
        if 'difficult' in pn and len(p) < 60: mode = 'diff'; continue
        if 'perspect' in pn and len(p) < 60: mode = 'persp'; continue
        if len(p) < 8: continue
        if mode == 'diff' and len(diffs) < 5 and len(p) > 15: diffs.append(p)
        elif mode == 'persp' and len(persps) < 5 and len(p) > 15: persps.append(p)
    return diffs, persps

# ═══════════════════════════════════════════════════════════
FILES_2024 = {
    0:  'Rapport Mensuel DNMF TOGO Janvier.docx',
    1:  'Rapport Mensuel DNMF TOGO F\u00e9vrier.docx',
    2:  'Rapport Mensuel DNMF TOGO Mars.docx',
    3:  'Rapport Mensuel DNMF TOGO Avril.docx',
    4:  'Rapport Mensuel DNMF TOGO MAI 2024.docx',
    5:  'Rapport Mensuel DNMF TOGO JUIN 2024.docx',
    6:  'Rapport Mensuel DNMF TOGO JUILLET 2024.docx',
    7:  'Rapport Mensuel DNMF TOGO AO\u00dbT 2024.docx',
    8:  'Rapport Mensuel DNMF TOGO septembre 2024.docx',
    9:  'Rapport Mensuel DNMF TOGO OCTOBRE 2024.docx',
    10: 'Rapport Mensuel DNMF TOGO NOVEMBRE 2024.docx',
    11: 'Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx',
}
FILES_2025 = {
    0:  'DAMF RAPPORT NATIONAL TOGO Jan 2025.docx',
    1:  'DAMF RAPPORT NATIONAL TOGO FEV 2025.docx',
    2:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE MARS 2025.docx',
    3:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE AVRIL 2025.docx',
    4:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE MAI 2025.docx',
    5:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUIN 2025.docx',
    6:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUILLET 2025.docx',
    7:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE AOUT 2025.docx',
    8:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE SEPTEMBRE 2025.docx',
    9:  "DNMF TOGO- RAPPORT MENSUEL MOIS D'OCTOBRE 2025.docx",
    10: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE NOVEMBRE 2025.docx',
    11: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE DECEMBRE 2025.docx',
}

result = {}
for year, files_map, folder, extractor in [
    (2024, FILES_2024, '2024', 'row_based'),
    (2025, FILES_2025, '2025', 'col_based'),
]:
    year_data = []
    diffs_g, persps_g = [], []
    last_doc = None

    for mi in range(12):
        entry = dict(
            month=MONTH_NAMES_DISPLAY[mi], monthIndex=mi, year=year,
            sem_assemblees=0, sem_hors=0, sem_total=0,
            assistance=0, sauves=0, ajoutes=0,
            invites=0, temoignages=0, pourcentage=0.0,
            predicateurs=0, pasteurs=0,
            miss_nationaux=0, miss_internationaux=0,
            eleves_inscrits=0, eleves_actuels=0,
            membres=0, assemblees_count=0, districts=0,
            predicateurs_utilises=0,
        )
        fname = files_map.get(mi)
        path = os.path.join(BASE, folder, fname) if fname else None
        if path and os.path.exists(path):
            print('  [{}/{}] {}'.format(folder, mi, fname[:55]))
            try:
                doc = docx.Document(path)
                if extractor == 'row_based':
                    extracted = extract_2024(doc, mi)
                else:
                    extracted = extract_2025_multi(doc, mi)
                entry.update(extracted)
                d, p = get_diffs_persps(doc)
                if d: diffs_g = d
                if p: persps_g = p
            except Exception as e:
                print('    ERROR: {}'.format(e))

        if entry['sem_total'] == 0:
            entry['sem_total'] = entry['sem_assemblees'] + entry['sem_hors']
        if entry['sauves'] > 0 and entry['ajoutes'] > 0:
            entry['pourcentage'] = round(entry['ajoutes'] / entry['sauves'] * 100, 1)
        year_data.append(entry)

    result[str(year)] = {'monthlyData': year_data, 'difficultes': diffs_g, 'perspectives': persps_g}

out = os.path.join(BASE, 'tlwm-dashboard', 'src', 'data', 'realData.json')
with open(out, 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print('\n=== VERIFICATION ===')
for year in ['2024', '2025']:
    print('\n' + year)
    for m in result[year]['monthlyData']:
        flag = '' if m['sauves'] > 0 or m['sem_total'] > 0 else '  <-- vide'
        print('  ' + m['month'].ljust(12) +
              ' sem=' + str(m['sem_total']).rjust(3) +
              ' assist=' + str(m['assistance']).rjust(5) +
              ' sauves=' + str(m['sauves']).rjust(4) +
              ' ajoutes=' + str(m['ajoutes']).rjust(3) +
              ' pred=' + str(m['predicateurs']).rjust(3) + flag)
print('\nOutput: ' + out)
