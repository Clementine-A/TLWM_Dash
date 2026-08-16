"""
Extraction complète et structurée de tous les fichiers DNMF TOGO
Génère un fichier JSON prêt pour le dashboard React
"""
import os
import re
import json
import docx

BASE = r"c:\Users\ATD\Desktop\cle\TLWM\dash"
PY = r"C:\Users\ATD\AppData\Local\Programs\Python\Python310\python.exe"

MONTH_MAP = {
    'janv': 0, 'jan': 0, 'janvier': 0,
    'fev': 1, 'fév': 1, 'février': 1, 'fevrier': 1,
    'mars': 2,
    'avril': 3, 'avr': 3,
    'mai': 4,
    'juin': 5,
    'juil': 6, 'juillet': 6,
    'août': 7, 'aout': 7, 'aoû': 7,
    'sept': 8, 'septembre': 8,
    'oct': 9, 'octobre': 9,
    'nov': 10, 'novembre': 10,
    'déc': 11, 'dec': 11, 'décembre': 11,
}

MONTH_NAMES_FR = [
    'Janvier','Février','Mars','Avril','Mai','Juin',
    'Juillet','Août','Septembre','Octobre','Novembre','Décembre'
]

def safe_int(s):
    if not s:
        return 0
    s = s.strip().replace(' ', '').replace('\xa0','').replace(',','').replace('.','')
    try:
        return int(s)
    except:
        return 0

def get_table_by_header(doc, keyword):
    """Returns first table whose first cell contains keyword"""
    kw = keyword.lower()
    for table in doc.tables:
        if table.rows:
            first = table.rows[0].cells[0].text.lower()
            if kw in first:
                return table
    return None

def get_col_idx(header_row, month_idx):
    """Find column index for a given month (0=Jan)"""
    for i, cell in enumerate(header_row.cells):
        txt = cell.text.strip().lower()
        norm = txt.replace('û','u').replace('é','e').replace('è','e').replace('ô','o').replace('â','a')
        if norm in MONTH_MAP and MONTH_MAP[norm] == month_idx:
            return i
    return -1

def extract_row_by_keyword(table, row_keyword, col_idx):
    """Find a row containing row_keyword, return value at col_idx"""
    kw = row_keyword.lower()
    for row in table.rows:
        label = row.cells[0].text.lower()
        if kw in label:
            if col_idx < len(row.cells):
                return safe_int(row.cells[col_idx].text)
    return 0

def extract_month_data_2025(doc, month_idx):
    """Extract data from 2025 multi-month cumulative tables"""
    data = {}

    # === SCE Table ===
    sce = get_table_by_header(doc, 'sce')
    if sce and len(sce.rows) > 1:
        hdr = sce.rows[1]
        col = get_col_idx(hdr, month_idx)
        if col >= 0:
            data['districts'] = extract_row_by_keyword(sce, 'districts', col)
            data['assemblees_count'] = extract_row_by_keyword(sce, 'assembl', col)
            data['membres'] = extract_row_by_keyword(sce, 'membres', col)

    # === Séminaires Table ===
    sem = get_table_by_header(doc, 'sémin') or get_table_by_header(doc, 'semin') or get_table_by_header(doc, 'sem')
    if sem and len(sem.rows) > 1:
        hdr = sem.rows[1]
        col = get_col_idx(hdr, month_idx)
        if col >= 0:
            data['sem_assemblees'] = extract_row_by_keyword(sem, 'assembl', col)
            data['sem_hors'] = extract_row_by_keyword(sem, 'hors', col)
            data['sem_total'] = extract_row_by_keyword(sem, 'total', col)
            data['assistance'] = extract_row_by_keyword(sem, 'assistance', col)
            data['invites'] = extract_row_by_keyword(sem, 'invit', col)
            data['sauves'] = extract_row_by_keyword(sem, 'sauv', col)
            data['ajoutes'] = extract_row_by_keyword(sem, 'ajout', col)
            data['predicateurs_utilises'] = extract_row_by_keyword(sem, 'pr\xe9dicateurs utilis', col)
            data['temoignages'] = extract_row_by_keyword(sem, 't\xe9moignage', col)

    # === SPGFM Table ===
    spg = get_table_by_header(doc, 'spgfm')
    if spg and len(spg.rows) > 1:
        hdr = spg.rows[1]
        col = get_col_idx(hdr, month_idx)
        if col >= 0:
            data['predicateurs'] = extract_row_by_keyword(spg, 'pr\xe9dicateurs', col)
            data['pasteurs'] = extract_row_by_keyword(spg, 'pasteurs', col)
            data['miss_nationaux'] = extract_row_by_keyword(spg, 'nationaux', col)
            data['miss_internationaux'] = extract_row_by_keyword(spg, 'internationaux', col)

    # === SFA (formation acadé) ===
    sfa = get_table_by_header(doc, 'sfa')
    if sfa and len(sfa.rows) > 1:
        hdr = sfa.rows[1]
        col = get_col_idx(hdr, month_idx)
        if col >= 0:
            data['eleves_inscrits'] = extract_row_by_keyword(sfa, 'inscrits', col)
            data['eleves_actuels'] = extract_row_by_keyword(sfa, 'actuels', col)

    return data

def extract_month_data_2024_simple(doc):
    """Extract data from 2024 single-month reports (simpler structure)"""
    data = {}
    full_text = '\n'.join(p.text for p in doc.paragraphs)

    def find_num(patterns, text):
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                try:
                    return int(m.group(1).replace(' ','').replace(',',''))
                except:
                    pass
        return 0

    # Try tables first
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            row_text = [c.text.strip() for c in row.cells]
            rows_text.append(row_text)

        for row in rows_text:
            label = row[0].lower() if row else ''
            val_str = row[1] if len(row) > 1 else ''

            if 's\xe9minaires' in label or 'seminaires' in label or 'nbre de sem' in label:
                if 'assembl' in label:
                    data['sem_assemblees'] = safe_int(val_str)
                elif 'hors' in label:
                    data['sem_hors'] = safe_int(val_str)
                elif 'total' in label or ('s\xe9minaires' in label and 'assembl' not in label and 'hors' not in label):
                    if not data.get('sem_total'):
                        data['sem_total'] = safe_int(val_str)
            if 'assistance' in label:
                if not data.get('assistance'):
                    data['assistance'] = safe_int(val_str)
            if 'sauv' in label:
                if not data.get('sauves'):
                    data['sauves'] = safe_int(val_str)
            if 'ajout' in label:
                if not data.get('ajoutes'):
                    data['ajoutes'] = safe_int(val_str)
            if 'pr\xe9dicateurs' in label and 'utilis' not in label and not data.get('predicateurs'):
                data['predicateurs'] = safe_int(val_str)
            if 'pasteurs' in label and not data.get('pasteurs'):
                data['pasteurs'] = safe_int(val_str)
            if 'missionnaire' in label and 'national' in label:
                data['miss_nationaux'] = safe_int(val_str)
            if 'missionnaire' in label and 'international' in label:
                data['miss_internationaux'] = safe_int(val_str)
            if '\xe9l\xe8ves' in label and 'inscrits' in label:
                data['eleves_inscrits'] = safe_int(val_str)
            if '\xe9l\xe8ves' in label and 'actuels' in label:
                data['eleves_actuels'] = safe_int(val_str)
            if 'membres' in label and not data.get('membres'):
                data['membres'] = safe_int(val_str)
            if 'assembl' in label and 'districts' not in label and not data.get('assemblees_count'):
                if safe_int(val_str) > 0:
                    data['assemblees_count'] = safe_int(val_str)
            if 'districts' in label and not data.get('districts'):
                data['districts'] = safe_int(val_str)

    return data

def get_difficulties_perspectives(doc):
    """Extract difficulties and perspectives from paragraphs"""
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    diffs = []
    persps = []
    mode = None
    for p in paragraphs:
        pl = p.lower()
        if 'difficult' in pl:
            mode = 'diff'
            continue
        if 'perspect' in pl:
            mode = 'persp'
            continue
        if len(p) > 10 and p not in ['.','',' ']:
            if mode == 'diff' and len(diffs) < 5:
                diffs.append(p)
            elif mode == 'persp' and len(persps) < 5:
                persps.append(p)
    return diffs, persps

# ───────────────────────────────────────────────────────────
# Process files
# ───────────────────────────────────────────────────────────

FILES_2024 = {
    0:  "Rapport Mensuel DNMF TOGO Janvier.docx",
    1:  "Rapport Mensuel DNMF TOGO Février.docx",
    2:  "Rapport Mensuel DNMF TOGO Mars.docx",
    3:  "Rapport Mensuel DNMF TOGO Avril.docx",
    4:  "Rapport Mensuel DNMF TOGO MAI 2024.docx",
    5:  "Rapport Mensuel DNMF TOGO JUIN 2024.docx",
    6:  "Rapport Mensuel DNMF TOGO JUILLET 2024.docx",
    7:  "Rapport Mensuel DNMF TOGO AOÛT 2024.docx",
    8:  "Rapport Mensuel DNMF TOGO septembre 2024.docx",
    9:  "Rapport Mensuel DNMF TOGO OCTOBRE 2024.docx",
    10: "Rapport Mensuel DNMF TOGO NOVEMBRE 2024.docx",
    11: "Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx",
}

FILES_2025 = {
    0:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE JANVIER 2025.docx",
    1:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE FEVRIER 2025.docx",
    2:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE MARS 2025.docx",
    3:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE AVRIL 2025.docx",
    4:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE MAI 2025.docx",
    5:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE JUIN 2025.docx",
    6:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE JUILLET 2025.docx",
    7:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE AOUT 2025.docx",
    8:  "DNMF TOGO- RAPPORT MENSUEL MOIS DE SEPTEMBRE 2025.docx",
    9:  "DNMF TOGO- RAPPORT MENSUEL MOIS D'OCTOBRE 2025.docx",
    10: "DNMF TOGO- RAPPORT MENSUEL MOIS DE NOVEMBRE 2025.docx",
    11: "DNMF TOGO- RAPPORT MENSUEL MOIS DE DECEMBRE 2025.docx",
}

out_path = os.path.join(BASE, "tlwm-dashboard", "src", "data", "realData.json")
if os.path.exists(out_path):
    with open(out_path, 'r', encoding='utf-8') as f:
        result = json.load(f)
else:
    result = {}

for year, files_map, folder, extractor_func in [
    (2025, FILES_2025, "2025", "cumulative"),
]:
    year_data = []
    diffs_year = []
    persps_year = []

    for month_idx in range(12):
        filename = files_map.get(month_idx)
        path = os.path.join(BASE, folder, filename) if filename else None

        entry = {
            "month": MONTH_NAMES_FR[month_idx],
            "monthIndex": month_idx,
            "year": year,
            "sem_assemblees": 0,
            "sem_hors": 0,
            "sem_total": 0,
            "assistance": 0,
            "sauves": 0,
            "ajoutes": 0,
            "invites": 0,
            "temoignages": 0,
            "predicateurs": 0,
            "pasteurs": 0,
            "miss_nationaux": 0,
            "miss_internationaux": 0,
            "eleves_inscrits": 0,
            "eleves_actuels": 0,
            "membres": 0,
            "assemblees_count": 0,
            "districts": 0,
        }

        if path and os.path.exists(path):
            print(f"Processing: {year}/{filename}")
            try:
                doc = docx.Document(path)
                if extractor_func == "cumulative":
                    extracted = extract_month_data_2025(doc, month_idx)
                else:
                    extracted = extract_month_data_2024_simple(doc)
                entry.update(extracted)

                # Get difficulties from last file of year (cumulative)
                if not diffs_year or month_idx == 11:
                    d, p = get_difficulties_perspectives(doc)
                    if d:
                        diffs_year = d
                    if p:
                        persps_year = p
            except Exception as e:
                print(f"  ERROR: {e}")
        else:
            print(f"Missing: {year}/{filename}")

        # Compute derived fields
        if entry['sem_total'] == 0:
            entry['sem_total'] = entry['sem_assemblees'] + entry['sem_hors']
        if entry['sauves'] > 0 and entry['ajoutes'] > 0:
            entry['pourcentage'] = round(entry['ajoutes'] / entry['sauves'] * 100, 1)
        else:
            entry['pourcentage'] = 0.0

        year_data.append(entry)

    result[str(year)] = {
        "monthlyData": year_data,
        "difficultes": diffs_year,
        "perspectives": persps_year,
    }

# Save JSON
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print(f"\n[OK] Données exportées vers: {out_path}")

# Print summary
for year in ['2025']:
    print(f"\n=== {year} ===")
    for m in result[year]['monthlyData']:
        print(f"  {m['month']:12s}: sem={m['sem_total']:3d} assist={m['assistance']:5d} sauves={m['sauves']:4d} ajoutes={m['ajoutes']:4d}")
