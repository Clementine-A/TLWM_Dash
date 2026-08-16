"""
Extraction finale CORRIGEE:
- 2024: lire DECEMBRE (contient toutes les données Jan->Déc) + fichiers individuels pour predicateurs
- 2025 Jan/Feb: lire FEV 2025 (contient Jan+Fev) depuis le grand tableau multi-colonne
- 2025 Mars+: même logique qu'avant
"""
import os, re, json
import unicodedata
import docx

BASE = r"c:\Users\ATD\Desktop\cle\TLWM\dash"

MONTH_NAMES_DISPLAY = [
    'Janvier','F\u00e9vrier','Mars','Avril','Mai','Juin',
    'Juillet','Ao\u00fbt','Septembre','Octobre','Novembre','D\u00e9cembre'
]

def norm(s):
    s = unicodedata.normalize('NFD', str(s).lower())
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn').strip()

MONTH_ID = {
    'janvier': 0, 'janv': 0, 'jan': 0,
    'fevrier': 1, 'fev': 1,
    'mars': 2,
    'avril': 3, 'avr': 3,
    'mai': 4,
    'juin': 5,
    'juillet': 6, 'juil': 6, 'juill': 6,
    'aout': 7, 'ao': 7,
    'septembre': 8, 'sept': 8, 'sep': 8,
    'octobre': 9, 'oct': 9,
    'novembre': 10, 'nov': 10,
    'decembre': 11, 'dec': 11,
}

def safe_int(s):
    m = re.search(r'\d[\d\s]*', str(s))
    if m:
        try:
            return int(m.group(0).replace(' ', ''))
        except:
            pass
    return 0

def match_month(cell_text):
    """Return month index from cell text, or None"""
    n = norm(cell_text)
    # Try exact
    if n in MONTH_ID:
        return MONTH_ID[n]
    # Try prefix match
    for k, v in sorted(MONTH_ID.items(), key=lambda x: -len(x[0])):
        if n.startswith(k) and len(k) >= 3:
            return v
    return None

# ─── Extract ALL months from a 2024-style cumulative file ─────
def extract_2024_all_months(doc):
    """
    Returns dict {month_idx: {sem_total, assistance, sauves, ajoutes, sem_assemblees, sem_hors}}
    from the cumulative tables (rows=months)
    """
    months = {}
    predicateurs_row0 = {}

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        r0 = norm(table.rows[0].cells[0].text)
        r1_text = table.rows[1].cells[0].text if len(table.rows) > 1 else ''
        r1 = norm(r1_text)

        # === Ressources humaines (Table 0): header row contains "predicateur" ===
        if 'predicateur' in r0 and len(table.rows) >= 2:
            row = table.rows[1]
            predicateurs_row0 = {
                'predicateurs': safe_int(row.cells[0].text) if len(row.cells) > 0 else 0,
                'eleves_inscrits': safe_int(row.cells[1].text) if len(row.cells) > 1 else 0,
                'miss_nationaux': safe_int(row.cells[2].text) if len(row.cells) > 2 else 0,
                'miss_internationaux': safe_int(row.cells[3].text) if len(row.cells) > 3 else 0,
            }
            continue

        # Cumulative tables have row1 cell0 == 'mois'
        if r1 != 'mois':
            continue

        section = r0
        for row in table.rows[2:]:
            if not row.cells:
                continue
            mi = match_month(row.cells[0].text)
            if mi is None:
                continue
            if mi not in months:
                months[mi] = {}

            if 'seminaire' in section or 'urgence' in section:
                months[mi]['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                months[mi]['assistance'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                months[mi]['sauves'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

            elif 'suivi' in section and 'sauve' in section:
                if 'sauves' not in months[mi]:
                    months[mi]['sauves'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                months[mi]['ajoutes'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0

            elif 'repartition' in section:
                if 'sem_total' not in months[mi]:
                    months[mi]['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                months[mi]['sem_assemblees'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                months[mi]['sem_hors'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

    return months, predicateurs_row0

# ─── 2025 multi-col format ───────────────────────────────────
def find_month_col(header_row, month_idx):
    for i, cell in enumerate(header_row.cells):
        tok = norm(cell.text).split()[0] if norm(cell.text).split() else ''
        if MONTH_ID.get(tok) == month_idx:
            return i
    return -1

def row_val(table, kw, col):
    for row in table.rows:
        if kw in norm(row.cells[0].text):
            return safe_int(row.cells[col].text) if col < len(row.cells) else 0
    return 0

def extract_2025_col_based(doc, month_idx):
    data = {}
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        h0 = norm(table.rows[0].cells[0].text)
        col = find_month_col(table.rows[1], month_idx)
        if col < 0:
            continue

        if 'sce' in h0 or 'etat' in h0:
            data['assemblees_count'] = row_val(table, 'assembl', col)
            data['membres'] = row_val(table, 'membres', col)
            data['districts'] = row_val(table, 'district', col)
        elif 'semin' in h0 or (h0.startswith('sem') and 'sem' in h0):
            data['sem_assemblees'] = row_val(table, 'assembl', col)
            data['sem_hors'] = row_val(table, 'hors', col)
            data['sem_total'] = row_val(table, 'total', col)
            data['assistance'] = row_val(table, 'assistance', col)
            data['sauves'] = row_val(table, 'sauv', col)
            data['ajoutes'] = row_val(table, 'ajout', col)
            data['invites'] = row_val(table, 'invit', col)
            data['temoignages'] = row_val(table, 'temoignage', col)
            data['predicateurs_utilises'] = row_val(table, 'predicateurs utilis', col)
        elif 'spgfm' in h0:
            data['predicateurs'] = row_val(table, 'nbre de pred', col) or row_val(table, 'predicateurs', col)
            data['pasteurs'] = row_val(table, 'pasteurs', col)
            data['miss_nationaux'] = row_val(table, 'national', col)
            data['miss_internationaux'] = row_val(table, 'international', col)
        elif 'sfa' in h0:
            data['eleves_inscrits'] = row_val(table, 'inscrits', col)
            data['eleves_actuels'] = row_val(table, 'actuels', col)
    return data

def get_diffs_persps(doc):
    paras = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]
    diffs, persps = [], []
    mode = None
    for p in paras:
        pn = norm(p)
        if 'difficult' in pn and len(p) < 60: mode = 'diff'; continue
        if 'perspect' in pn and len(p) < 60: mode = 'persp'; continue
        if len(p) < 8: continue
        if mode == 'diff' and len(diffs) < 6 and len(p) > 15: diffs.append(p)
        elif mode == 'persp' and len(persps) < 6 and len(p) > 15: persps.append(p)
    return diffs, persps

# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
result = {}

# ─── 2024: use DECEMBRE file (has all 12 months) ────────────
print('=== 2024 (reading from DECEMBRE + individual pred) ===')
dec_path = os.path.join(BASE, '2024', 'Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx')
doc_dec = docx.Document(dec_path)
all_2024_months, pred_dec = extract_2024_all_months(doc_dec)
print('  Months found in Décembre file:', sorted(all_2024_months.keys()))

# Read predicateurs per individual month files
FILES_2024 = {
    0: 'Rapport Mensuel DNMF TOGO Janvier.docx',
    1: 'Rapport Mensuel DNMF TOGO F\u00e9vrier.docx',
    2: 'Rapport Mensuel DNMF TOGO Mars.docx',
    3: 'Rapport Mensuel DNMF TOGO Avril.docx',
    4: 'Rapport Mensuel DNMF TOGO MAI 2024.docx',
    5: 'Rapport Mensuel DNMF TOGO JUIN 2024.docx',
    6: 'Rapport Mensuel DNMF TOGO JUILLET 2024.docx',
    7: 'Rapport Mensuel DNMF TOGO AO\u00dbT 2024.docx',
    8: 'Rapport Mensuel DNMF TOGO septembre 2024.docx',
    9: 'Rapport Mensuel DNMF TOGO OCTOBRE 2024.docx',
    10: 'Rapport Mensuel DNMF TOGO NOVEMBRE 2024.docx',
    11: 'Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx',
}
pred_per_month = {}
for mi, fname in FILES_2024.items():
    path = os.path.join(BASE, '2024', fname)
    if os.path.exists(path):
        try:
            doc_m = docx.Document(path)
            _, pr = extract_2024_all_months(doc_m)
            pred_per_month[mi] = pr
        except:
            pass

year_data_2024 = []
diffs_2024, persps_2024 = get_diffs_persps(doc_dec)
for mi in range(12):
    entry = dict(
        month=MONTH_NAMES_DISPLAY[mi], monthIndex=mi, year=2024,
        sem_assemblees=0, sem_hors=0, sem_total=0,
        assistance=0, sauves=0, ajoutes=0,
        invites=0, temoignages=0, pourcentage=0.0,
        predicateurs=0, pasteurs=0,
        miss_nationaux=0, miss_internationaux=0,
        eleves_inscrits=0, eleves_actuels=0,
        membres=0, assemblees_count=0, districts=0,
        predicateurs_utilises=0,
    )
    if mi in all_2024_months:
        entry.update(all_2024_months[mi])
    pr = pred_per_month.get(mi, pred_dec)
    entry.update({k: v for k, v in pr.items() if k not in entry or entry[k] == 0})

    if entry['sem_total'] == 0:
        entry['sem_total'] = entry['sem_assemblees'] + entry['sem_hors']
    if entry['sauves'] > 0 and entry['ajoutes'] > 0:
        entry['pourcentage'] = round(entry['ajoutes'] / entry['sauves'] * 100, 1)
    year_data_2024.append(entry)
result['2024'] = {'monthlyData': year_data_2024, 'difficultes': diffs_2024, 'perspectives': persps_2024}

# ─── 2025 ────────────────────────────────────────────────────
print('\n=== 2025 ===')
FILES_2025 = {
    0:  'DAMF RAPPORT NATIONAL TOGO Jan 2025.docx',
    1:  'DAMF RAPPORT NATIONAL TOGO FEV 2025.docx',
    2:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE MARS 2025.docx',
    3:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE AVRIL 2025.docx',
    4:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE MAI 2025.docx',
    5:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUIN 2025.docx',
    6:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUILLET 2025.docx',
    7:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE AOUT 2025.docx',
    8:  'DNMF TOGO- RAPPORT MENSUEL MOIS DE SEPTEMBRE 2025.docx',
    9:  "DNMF TOGO- RAPPORT MENSUEL MOIS D'OCTOBRE 2025.docx",
    10: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE NOVEMBRE 2025.docx',
    11: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE DECEMBRE 2025.docx',
}
year_data_2025 = []
diffs_2025, persps_2025 = [], []
for mi in range(12):
    entry = dict(
        month=MONTH_NAMES_DISPLAY[mi], monthIndex=mi, year=2025,
        sem_assemblees=0, sem_hors=0, sem_total=0,
        assistance=0, sauves=0, ajoutes=0,
        invites=0, temoignages=0, pourcentage=0.0,
        predicateurs=0, pasteurs=0,
        miss_nationaux=0, miss_internationaux=0,
        eleves_inscrits=0, eleves_actuels=0,
        membres=0, assemblees_count=0, districts=0,
        predicateurs_utilises=0,
    )
    fname = FILES_2025.get(mi)
    path = os.path.join(BASE, '2025', fname) if fname else None
    if path and os.path.exists(path):
        print('  [{}] {}'.format(mi, fname[:55]))
        try:
            doc = docx.Document(path)
            extracted = extract_2025_col_based(doc, mi)
            entry.update(extracted)
            d, p = get_diffs_persps(doc)
            if d: diffs_2025 = d
            if p: persps_2025 = p
        except Exception as e:
            print('    ERROR: {}'.format(e))

    if entry['sem_total'] == 0:
        entry['sem_total'] = entry['sem_assemblees'] + entry['sem_hors']
    if entry['sauves'] > 0 and entry['ajoutes'] > 0:
        entry['pourcentage'] = round(entry['ajoutes'] / entry['sauves'] * 100, 1)
    year_data_2025.append(entry)

result['2025'] = {'monthlyData': year_data_2025, 'difficultes': diffs_2025, 'perspectives': persps_2025}

# ─── Save ────────────────────────────────────────────────────
out = os.path.join(BASE, 'tlwm-dashboard', 'src', 'data', 'realData.json')
with open(out, 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print('\n=== FINAL VERIFICATION ===')
for year in ['2024', '2025']:
    print('\n' + year)
    total_sem = total_assist = total_sauves = total_ajoutes = 0
    for m in result[year]['monthlyData']:
        flag = '' if (m['sauves'] > 0 or m['sem_total'] > 0) else '  <-- VIDE'
        print('  ' + m['month'].ljust(12) +
              ' sem=' + str(m['sem_total']).rjust(3) +
              ' assist=' + str(m['assistance']).rjust(5) +
              ' sauves=' + str(m['sauves']).rjust(4) +
              ' ajoutes=' + str(m['ajoutes']).rjust(3) + flag)
        total_sem += m['sem_total']
        total_assist += m['assistance']
        total_sauves += m['sauves']
        total_ajoutes += m['ajoutes']
    print('  TOTAL: sem={} assist={} sauves={} ajoutes={}'.format(
        total_sem, total_assist, total_sauves, total_ajoutes))
print('\nOutput -> ' + out)
