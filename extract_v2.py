"""
Extracteur optimise pour les fichiers DNMF TOGO
- 2024: tableaux mensuels cumulatifs (rows=mois, cols=indicateurs)
- 2025 mars+: grand tableau annuel (rows=indicateurs, cols=mois)
- 2025 jan/fev: grand tableau complexe
"""
import os, re, json
import docx

BASE = r"c:\Users\ATD\Desktop\cle\TLWM\dash"

MONTH_NAMES_FR = [
    'Janvier','Fevrier','Mars','Avril','Mai','Juin',
    'Juillet','Aout','Septembre','Octobre','Novembre','Decembre'
]
MONTH_NAMES_DISPLAY = [
    'Janvier','F\u00e9vrier','Mars','Avril','Mai','Juin',
    'Juillet','Ao\u00fbt','Septembre','Octobre','Novembre','D\u00e9cembre'
]

MONTH_ROW_MAP = {
    'janvier': 0, 'janv': 0, 'jan': 0,
    'fevrier': 1, 'fev': 1, 'f': 1,
    'mars': 2,
    'avril': 3, 'avr': 3,
    'mai': 4,
    'juin': 5,
    'juillet': 6, 'juil': 6,
    'aout': 7, 'ao': 7,
    'septembre': 8, 'sept': 8, 'sep': 8,
    'octobre': 9, 'oct': 9,
    'novembre': 10, 'nov': 10,
    'decembre': 11, 'dec': 11,
}

def normalize(s):
    import unicodedata
    s = unicodedata.normalize('NFD', s.lower())
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn')

def safe_int(s):
    if not s:
        return 0
    s = str(s).strip().replace(' ','').replace('\xa0','').replace(',','').replace('.','').replace('-','')
    try:
        return int(s)
    except:
        return 0

# ─────────────────────────────────────────────────────────────
# 2024 format: rows = months, cols = values
# ─────────────────────────────────────────────────────────────
def extract_2024(doc, month_idx):
    """
    2024 files have cumulative tables:
    Table 1: Séminaires/Urgences (rows=months, cols=Total|Assistance|Declares sauves)
    Table 2: Suivi des sauvés (rows=months, cols=Declares sauves|Ajoutes)
    Table 3: Répartition (rows=months, cols=Total|Assemblees|Hors assemblees)
    Table 0: Ressources humaines (1 row)
    """
    data = {}

    for table in doc.tables:
        if not table.rows:
            continue
        # Identify by first cell of header row
        header0 = normalize(table.rows[0].cells[0].text)
        header1 = normalize(table.rows[1].cells[0].text) if len(table.rows) > 1 else ''

        # === Ressources humaines (Table 0) ===
        if 'predicateurs' in header0 and len(table.rows) >= 2:
            row = table.rows[1]
            data['predicateurs'] = safe_int(row.cells[0].text)
            data['eleves_inscrits'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
            data['miss_nationaux'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
            data['miss_internationaux'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

        # === Séminaires (Table 1) ===
        elif ('seminaire' in header0 or 'urgence' in header0) and header1 == 'mois':
            for row in table.rows[2:]:
                month_cell = normalize(row.cells[0].text)
                mi = MONTH_ROW_MAP.get(month_cell)
                if mi == month_idx:
                    data['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                    data['assistance'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                    data['sauves'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

        # === Suivi des sauvés (Table 2) ===
        elif 'suivi' in header0 and header1 == 'mois':
            for row in table.rows[2:]:
                month_cell = normalize(row.cells[0].text)
                mi = MONTH_ROW_MAP.get(month_cell)
                if mi == month_idx:
                    data['sauves'] = data.get('sauves') or safe_int(row.cells[1].text)
                    data['ajoutes'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0

        # === Répartition (Table 3) ===
        elif 'repartition' in header0 and header1 == 'mois':
            for row in table.rows[2:]:
                month_cell = normalize(row.cells[0].text)
                mi = MONTH_ROW_MAP.get(month_cell)
                if mi == month_idx:
                    data['sem_total'] = data.get('sem_total') or safe_int(row.cells[1].text)
                    data['sem_assemblees'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                    data['sem_hors'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

        # === Assemblées ===
        elif 'assemblee' in header0 and 'nombre' in header0:
            for row in table.rows[1:]:
                val = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                if val > 0 and 'assemblee' not in data:
                    data['assemblees_count'] = val

    return data

# ─────────────────────────────────────────────────────────────
# 2025 format (Mars+): rows=indicateurs, cols=Jan..mois_courant
# ─────────────────────────────────────────────────────────────
def get_col_2025(header_row, month_idx):
    for i, cell in enumerate(header_row.cells):
        n = normalize(cell.text.strip())
        # strip to first word
        n = n.split()[0] if n.split() else n
        if MONTH_ROW_MAP.get(n) == month_idx:
            return i
    return -1

def extract_row_val(table, kw, col):
    for row in table.rows:
        lbl = normalize(row.cells[0].text)
        if kw in lbl:
            if col >= 0 and col < len(row.cells):
                return safe_int(row.cells[col].text)
    return 0

def extract_2025_multi(doc, month_idx):
    data = {}
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header0 = normalize(table.rows[0].cells[0].text)
        hdr = table.rows[1] if len(table.rows) > 1 else table.rows[0]

        col = get_col_2025(hdr, month_idx)
        if col < 0:
            continue

        if 'sce' in header0 or 'etat' in header0:
            data['assemblees_count'] = extract_row_val(table, 'assembl', col)
            data['membres'] = extract_row_val(table, 'membres', col)
            data['districts'] = extract_row_val(table, 'district', col)

        elif 'semin' in header0 or 'sem' in header0:
            data['sem_assemblees'] = extract_row_val(table, 'assembl', col)
            data['sem_hors'] = extract_row_val(table, 'hors', col)
            data['sem_total'] = extract_row_val(table, 'total', col)
            data['assistance'] = extract_row_val(table, 'assistance', col)
            data['sauves'] = extract_row_val(table, 'sauv', col)
            data['ajoutes'] = extract_row_val(table, 'ajout', col)
            data['invites'] = extract_row_val(table, 'invit', col)
            data['temoignages'] = extract_row_val(table, 'temoignage', col)
            data['predicateurs_utilises'] = extract_row_val(table, 'predicateurs utilis', col)

        elif 'spgfm' in header0:
            data['predicateurs'] = extract_row_val(table, 'predicateurs', col)
            data['pasteurs'] = extract_row_val(table, 'pasteurs', col)
            data['miss_nationaux'] = extract_row_val(table, 'national', col)
            data['miss_internationaux'] = extract_row_val(table, 'international', col)

        elif 'sfa' in header0:
            data['eleves_inscrits'] = extract_row_val(table, 'inscrits', col)
            data['eleves_actuels'] = extract_row_val(table, 'actuels', col)

    return data

# ─────────────────────────────────────────────────────────────
# 2025 Jan/Fev: mega-table format (big flat table)
# ─────────────────────────────────────────────────────────────
def extract_2025_jan_fev(doc, month_idx):
    """Jan/Fev 2025 have a different big single-table format"""
    data = {}
    text = '\n'.join(p.text for p in doc.paragraphs)
    # Try to find key numbers from paragraph text
    patterns = {
        'sem_total': [r'(?:nombre|nbre).*?s[eé]minaires?\s*[:\-]?\s*(\d+)', r'total.*?s[eé]minaires?\s*[:\-]?\s*(\d+)'],
        'assistance': [r'assistance\s*[:\-]?\s*(\d+)'],
        'sauves': [r'sauv[eé]s?\s*[:\-]?\s*(\d+)', r'd[eé]clar[eé]s?\s*sauv[eé]s?\s*[:\-]?\s*(\d+)'],
        'ajoutes': [r'ajout[eé]s?\s*[:\-]?\s*(\d+)'],
        'predicateurs': [r'pr[eé]dicateurs?\s*[:\-]?\s*(\d+)'],
    }
    for key, pats in patterns.items():
        for pat in pats:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                data[key] = safe_int(m.group(1))
                break

    # Try tables too
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cn = normalize(cell.text)
                if ci + 1 < len(row.cells):
                    val = safe_int(row.cells[ci+1].text)
                    if 'predicateurs' in cn and not 'eleves' in cn and not data.get('predicateurs') and val > 0:
                        data['predicateurs'] = val
                    if 'assistance' in cn and not data.get('assistance') and val > 0:
                        data['assistance'] = val
                    if 'sauves' in cn and not data.get('sauves') and val > 0:
                        data['sauves'] = val
                    if 'ajoutes' in cn and not data.get('ajoutes') and val > 0:
                        data['ajoutes'] = val
    return data

def get_diffs_persps(doc):
    paras = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]
    diffs, persps = [], []
    mode = None
    for p in paras:
        pn = normalize(p)
        if 'difficult' in pn and len(p) < 60:
            mode = 'diff'
            continue
        if 'perspect' in pn and len(p) < 60:
            mode = 'persp'
            continue
        if p in ['.','',' '] or len(p) < 5:
            continue
        if mode == 'diff' and len(diffs) < 6 and len(p) > 15:
            diffs.append(p)
        elif mode == 'persp' and len(persps) < 6 and len(p) > 15:
            persps.append(p)
    return diffs, persps

# ═══════════════════════════════════════════════════════════
# File manifests
# ═══════════════════════════════════════════════════════════
FILES_2024 = {
    0: 'Rapport Mensuel DNMF TOGO Janvier.docx',
    1: 'Rapport Mensuel DNMF TOGO F\u00e9vrier.docx',
    2: 'Rapport Mensuel DNMF TOGO Mars.docx',
    3: 'Rapport Mensuel DNMF TOGO Avril.docx',
    4: 'Rapport Mensuel DNMF TOGO MAI 2024.docx',
    5: 'Rapport Mensuel DNMF TOGO JUIN 2024.docx',
    6: 'Rapport Mensuel DNMF TOGO JUILLET 2024.docx',
    7: 'Rapport Mensuel DNMF TOGO AO\u00dbT 2024.docx',
    8: 'Rapport Mensuel DNMF TOGO septembre 2024.docx',
    9: 'Rapport Mensuel DNMF TOGO OCTOBRE 2024.docx',
    10: 'Rapport Mensuel DNMF TOGO NOVEMBRE 2024.docx',
    11: 'Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx',
}

FILES_2025 = {
    0: 'DAMF RAPPORT NATIONAL TOGO Jan 2025.docx',
    1: 'DAMF RAPPORT NATIONAL TOGO FEV 2025.docx',
    2: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE MARS 2025.docx',
    3: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE AVRIL 2025.docx',
    4: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE MAI 2025.docx',
    5: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUIN 2025.docx',
    6: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUILLET 2025.docx',
    7: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE AOUT 2025.docx',
    8: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE SEPTEMBRE 2025.docx',
    9: "DNMF TOGO- RAPPORT MENSUEL MOIS D'OCTOBRE 2025.docx",
    10: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE NOVEMBRE 2025.docx',
    11: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE DECEMBRE 2025.docx',
}

result = {}

for year, files_map, folder in [(2024, FILES_2024, '2024'), (2025, FILES_2025, '2025')]:
    year_data = []
    diffs_global, persps_global = [], []

    for month_idx in range(12):
        fname = files_map.get(month_idx)
        path = os.path.join(BASE, folder, fname) if fname else None

        entry = dict(
            month=MONTH_NAMES_DISPLAY[month_idx],
            monthIndex=month_idx,
            year=year,
            sem_assemblees=0, sem_hors=0, sem_total=0,
            assistance=0, sauves=0, ajoutes=0,
            invites=0, temoignages=0, pourcentage=0.0,
            predicateurs=0, pasteurs=0,
            miss_nationaux=0, miss_internationaux=0,
            eleves_inscrits=0, eleves_actuels=0,
            membres=0, assemblees_count=0, districts=0,
            predicateurs_utilises=0,
        )

        if path and os.path.exists(path):
            print('Processing: {}/{}'.format(folder, fname[:50]))
            try:
                doc = docx.Document(path)
                if year == 2024:
                    extracted = extract_2024(doc, month_idx)
                elif year == 2025 and month_idx <= 1:
                    extracted = extract_2025_jan_fev(doc, month_idx)
                else:
                    extracted = extract_2025_multi(doc, month_idx)
                entry.update(extracted)

                d, p = get_diffs_persps(doc)
                if d:
                    diffs_global = d
                if p:
                    persps_global = p
            except Exception as e:
                print('  ERROR in {}: {}'.format(fname, e))
        else:
            print('Missing: {}/{}'.format(folder, fname))

        if entry['sem_total'] == 0:
            entry['sem_total'] = entry['sem_assemblees'] + entry['sem_hors']
        if entry['sauves'] > 0 and entry['ajoutes'] > 0:
            entry['pourcentage'] = round(entry['ajoutes'] / entry['sauves'] * 100, 1)

        year_data.append(entry)

    result[str(year)] = {
        'monthlyData': year_data,
        'difficultes': diffs_global,
        'perspectives': persps_global,
    }

out_path = os.path.join(BASE, 'tlwm-dashboard', 'src', 'data', 'realData.json')
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print('\nDone -> ' + out_path)
for year in ['2024', '2025']:
    print('\n=== ' + year + ' ===')
    for m in result[year]['monthlyData']:
        print('  ' + m['month'].ljust(12) +
              ' sem=' + str(m['sem_total']).rjust(3) +
              ' assist=' + str(m['assistance']).rjust(5) +
              ' sauves=' + str(m['sauves']).rjust(4) +
              ' ajoutes=' + str(m['ajoutes']).rjust(3))
