"""
EXTRACTEUR DEFINITIF — DNMF TOGO 2024 & 2025
Correction: la table header pour 2024 est sur ROW[1] (sous-entete), pas ROW[0]
"""
import os, re, json, unicodedata
import docx

BASE = r"c:\Users\ATD\Desktop\cle\TLWM\dash"

MONTH_NAMES_DISPLAY = [
    'Janvier','F\xe9vrier','Mars','Avril','Mai','Juin',
    'Juillet','Ao\xfbt','Septembre','Octobre','Novembre','D\xe9cembre'
]

def norm(s):
    s = unicodedata.normalize('NFD', str(s).lower())
    return ''.join(c for c in s if unicodedata.category(c) != 'Mn').strip()

MONTH_ID = {
    'janvier': 0, 'janv': 0, 'jan': 0,
    'fevrier': 1, 'fev': 1,
    'mars': 2,
    'avril': 3, 'avr': 3,
    'mai': 4,
    'juin': 5,
    'juillet': 6, 'juil': 6, 'juill': 6,
    'aout': 7,
    'septembre': 8, 'sept': 8, 'sep': 8,
    'octobre': 9, 'oct': 9,
    'novembre': 10, 'nov': 10,
    'decembre': 11, 'dec': 11,
}

def safe_int(s):
    m = re.search(r'\d[\d\s]*', str(s))
    if m:
        try: return int(m.group(0).replace(' ', ''))
        except: pass
    return 0

def match_month(cell_text):
    n = norm(cell_text)
    if n in MONTH_ID: return MONTH_ID[n]
    for k in sorted(MONTH_ID, key=len, reverse=True):
        if n.startswith(k) and len(k) >= 3: return MONTH_ID[k]
    return None

def extract_2024_all(doc):
    """Extract all months data from 2024 cumulative tables.
    Row[0].cell[0] = 'Mois' (section header)
    Row[1] = column subheaders: 'Mois | Total | Assistance | Declares sauves' etc.
    Row[2+] = data rows by month
    """
    months = {}
    pred_data = {}

    for t_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2: continue

        r0n = norm(table.rows[0].cells[0].text)

        # === Table 0: RH (no mois row) ===
        if 'predicateur' in r0n and 'nombre' in r0n:
            row = table.rows[1]
            pred_data = {
                'predicateurs': safe_int(row.cells[0].text) if len(row.cells) > 0 else 0,
                'eleves_inscrits': safe_int(row.cells[1].text) if len(row.cells) > 1 else 0,
                'miss_nationaux': safe_int(row.cells[2].text) if len(row.cells) > 2 else 0,
                'miss_internationaux': safe_int(row.cells[3].text) if len(row.cells) > 3 else 0,
                'predicateurs_utilises': safe_int(row.cells[4].text) if len(row.cells) > 4 else 0,
            }
            continue

        # Tables with first col = 'Mois'  -> cumulative monthly tables
        if r0n != 'mois': continue

        # Identify type by row[1] subheader cells
        if len(table.rows) < 3: continue
        sub = [norm(c.text) for c in table.rows[1].cells[:4]]

        # Data rows start at row 2
        for row in table.rows[2:]:
            if not row.cells: continue
            mi = match_month(row.cells[0].text)
            if mi is None: continue
            if mi not in months: months[mi] = {}

            # Séminaires / Urgences: cols = Mois|Total|Assistance|Declares sauves
            if any('total' in s for s in sub) and any('assistance' in s for s in sub):
                months[mi]['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                months[mi]['assistance'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                months[mi]['sauves'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

            # Suivi des sauvés: cols = Mois|Declares sauves|Ajoutes|%
            elif any('ajout' in s for s in sub) and any('sauve' in s or 'declare' in s for s in sub):
                if 'sauves' not in months[mi]:
                    months[mi]['sauves'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                months[mi]['ajoutes'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0

            # Répartition: cols = Mois|Total|Assemblees|Hors assemblees
            elif any('assemblee' in s or 'organis' in s for s in sub):
                if 'sem_total' not in months[mi]:
                    months[mi]['sem_total'] = safe_int(row.cells[1].text) if len(row.cells) > 1 else 0
                months[mi]['sem_assemblees'] = safe_int(row.cells[2].text) if len(row.cells) > 2 else 0
                months[mi]['sem_hors'] = safe_int(row.cells[3].text) if len(row.cells) > 3 else 0

    return months, pred_data

# ─── 2025 multi-col format ───────────────────────────────────
def find_col(hdr_row, month_idx):
    for i, cell in enumerate(hdr_row.cells):
        tok = norm(cell.text).split()
        tok = tok[0] if tok else ''
        if MONTH_ID.get(tok) == month_idx: return i
    return -1

def row_val(table, kw, col):
    for row in table.rows:
        if kw in norm(row.cells[0].text):
            return safe_int(row.cells[col].text) if col < len(row.cells) else 0
    return 0

def extract_2025(doc, month_idx):
    data = {}
    for table in doc.tables:
        if len(table.rows) < 2: continue
        h0 = norm(table.rows[0].cells[0].text)
        col = find_col(table.rows[1], month_idx)
        if col < 0: continue

        if 'sce' in h0 or 'etat' in h0:
            data['assemblees_count'] = row_val(table, 'assembl', col)
            data['membres'] = row_val(table, 'membres', col)
            data['districts'] = row_val(table, 'district', col)
        elif 'semin' in h0 or (h0.startswith('sem') and len(h0) < 20):
            data['sem_assemblees'] = row_val(table, 'assembl', col)
            data['sem_hors'] = row_val(table, 'hors', col)
            data['sem_total'] = row_val(table, 'total', col)
            data['assistance'] = row_val(table, 'assistance', col)
            data['sauves'] = row_val(table, 'sauv', col)
            data['ajoutes'] = row_val(table, 'ajout', col)
            data['invites'] = row_val(table, 'invit', col)
            data['temoignages'] = row_val(table, 'temoignage', col)
            data['predicateurs_utilises'] = row_val(table, 'predicateurs utilis', col)
        elif 'spgfm' in h0:
            data['predicateurs'] = row_val(table, 'predicateurs', col)
            data['pasteurs'] = row_val(table, 'pasteurs', col)
            data['miss_nationaux'] = row_val(table, 'national', col)
            data['miss_internationaux'] = row_val(table, 'international', col)
        elif 'sfa' in h0:
            data['eleves_inscrits'] = row_val(table, 'inscrits', col)
            data['eleves_actuels'] = row_val(table, 'actuels', col)
    return data

def get_diffs_persps(doc):
    paras = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]
    diffs, persps = [], []
    mode = None
    for p in paras:
        pn = norm(p)
        if 'difficult' in pn and len(p) < 60: mode = 'diff'; continue
        if 'perspect' in pn and len(p) < 60: mode = 'persp'; continue
        if len(p) < 8: continue
        if mode == 'diff' and len(diffs) < 6 and len(p) > 15: diffs.append(p)
        elif mode == 'persp' and len(persps) < 6 and len(p) > 15: persps.append(p)
    return diffs, persps

def blank_entry(mi, year):
    return dict(
        month=MONTH_NAMES_DISPLAY[mi], monthIndex=mi, year=year,
        sem_assemblees=0, sem_hors=0, sem_total=0,
        assistance=0, sauves=0, ajoutes=0,
        invites=0, temoignages=0, pourcentage=0.0,
        predicateurs=0, pasteurs=0,
        miss_nationaux=0, miss_internationaux=0,
        eleves_inscrits=0, eleves_actuels=0,
        membres=0, assemblees_count=0, districts=0,
        predicateurs_utilises=0,
    )

FILES_2024 = {
    0: 'Rapport Mensuel DNMF TOGO Janvier.docx',
    1: 'Rapport Mensuel DNMF TOGO F\xe9vrier.docx',
    2: 'Rapport Mensuel DNMF TOGO Mars.docx',
    3: 'Rapport Mensuel DNMF TOGO Avril.docx',
    4: 'Rapport Mensuel DNMF TOGO MAI 2024.docx',
    5: 'Rapport Mensuel DNMF TOGO JUIN 2024.docx',
    6: 'Rapport Mensuel DNMF TOGO JUILLET 2024.docx',
    7: 'Rapport Mensuel DNMF TOGO AO\xdbT 2024.docx',
    8: 'Rapport Mensuel DNMF TOGO septembre 2024.docx',
    9: 'Rapport Mensuel DNMF TOGO OCTOBRE 2024.docx',
    10: 'Rapport Mensuel DNMF TOGO NOVEMBRE 2024.docx',
    11: 'Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx',
}
FILES_2025 = {
    0: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE JANVIER 2025.docx',
    1: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE FEVRIER 2025.docx',
    2: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE MARS 2025.docx',
    3: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE AVRIL 2025.docx',
    4: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE MAI 2025.docx',
    5: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUIN 2025.docx',
    6: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE JUILLET 2025.docx',
    7: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE AOUT 2025.docx',
    8: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE SEPTEMBRE 2025.docx',
    9: "DNMF TOGO- RAPPORT MENSUEL MOIS D'OCTOBRE 2025.docx",
    10: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE NOVEMBRE 2025.docx',
    11: 'DNMF TOGO- RAPPORT MENSUEL MOIS DE DECEMBRE 2025.docx',
}

result = {}

# ────── 2024 ──────
print('=== Processing 2024 ===')
dec_path = os.path.join(BASE, '2024', 'Rapport Mensuel DNMF TOGO DECEMBRE 2024.docx')
doc_dec = docx.Document(dec_path)
all_months_2024, pred_dec = extract_2024_all(doc_dec)
diffs_2024, persps_2024 = get_diffs_persps(doc_dec)
print('  Months in DEC file:', sorted(all_months_2024.keys()))
for mi in sorted(all_months_2024.keys())[:3]:
    print('   mi={}: {}'.format(mi, all_months_2024[mi]))

# Per month predicateurs from individual files
pred_individual = {}
for mi, fname in FILES_2024.items():
    path = os.path.join(BASE, '2024', fname)
    if os.path.exists(path):
        try:
            _, pr = extract_2024_all(docx.Document(path))
            if pr: pred_individual[mi] = pr
        except: pass

year_data_2024 = []
for mi in range(12):
    e = blank_entry(mi, 2024)
    if mi in all_months_2024:
        e.update(all_months_2024[mi])
    pr = pred_individual.get(mi, pred_dec)
    for k, v in pr.items():
        if e.get(k, 0) == 0 and v: e[k] = v
    if e['sem_total'] == 0:
        e['sem_total'] = e['sem_assemblees'] + e['sem_hors']
    if e['sauves'] > 0 and e['ajoutes'] > 0:
        e['pourcentage'] = round(e['ajoutes'] / e['sauves'] * 100, 1)
    year_data_2024.append(e)
result['2024'] = {'monthlyData': year_data_2024, 'difficultes': diffs_2024, 'perspectives': persps_2024}

# ────── 2025 ──────
print('\n=== Processing 2025 ===')
year_data_2025 = []
diffs_2025, persps_2025 = [], []
for mi in range(12):
    e = blank_entry(mi, 2025)
    fname = FILES_2025.get(mi)
    path = os.path.join(BASE, '2025', fname) if fname else None
    if path and os.path.exists(path):
        try:
            doc = docx.Document(path)
            extracted = extract_2025(doc, mi)
            e.update(extracted)
            d, p = get_diffs_persps(doc)
            if d: diffs_2025 = d
            if p: persps_2025 = p
        except Exception as ex:
            print('  ERROR mi={}: {}'.format(mi, ex))
    if e['sem_total'] == 0:
        e['sem_total'] = e['sem_assemblees'] + e['sem_hors']
    if e['sauves'] > 0 and e['ajoutes'] > 0:
        e['pourcentage'] = round(e['ajoutes'] / e['sauves'] * 100, 1)
    year_data_2025.append(e)
result['2025'] = {'monthlyData': year_data_2025, 'difficultes': diffs_2025, 'perspectives': persps_2025}

out = os.path.join(BASE, 'tlwm-dashboard', 'src', 'data', 'realData.json')
with open(out, 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print('\n=== RESULTS ===')
for year in ['2024', '2025']:
    print('\n' + year)
    ts, ta, tsv, taj = 0, 0, 0, 0
    for m in result[year]['monthlyData']:
        flag = '  !!VIDE!!' if (m['sauves'] == 0 and m['sem_total'] == 0) else ''
        print('  ' + m['month'].ljust(12) +
              ' sem=' + str(m['sem_total']).rjust(3) +
              ' assist=' + str(m['assistance']).rjust(5) +
              ' sauves=' + str(m['sauves']).rjust(4) +
              ' ajoutes=' + str(m['ajoutes']).rjust(3) +
              ' pred=' + str(m['predicateurs']).rjust(3) + flag)
        ts += m['sem_total']; ta += m['assistance']
        tsv += m['sauves']; taj += m['ajoutes']
    print('  --- TOTAL: sem={} assit={} sauves={} ajoutes={}'.format(ts, ta, tsv, taj))
print('\nSaved -> ' + out)
